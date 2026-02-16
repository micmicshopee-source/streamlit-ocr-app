# -*- coding: utf-8 -*-
"""
AI PDF → Word 高品質版面還原 Prototype
使用 Gemini Vision 解析 PDF 每頁結構（標題、段落、清單、表格、圖片），
再用 python-docx 重建可編輯 Word（樣式、字體、圖片皆可編輯）。

使用方式：
  1. 設定 GEMINI_API_KEY 環境變數，或放在 .streamlit/secrets.toml 的 GEMINI_API_KEY
  2. 執行：python ai_pdf_to_word_prototype.py <PDF路徑> [輸出.docx路徑]
  3. 若未給輸出路徑，預設為 <PDF檔名>_ai_layout.docx

依賴：pip install pdf2image python-docx Pillow pymupdf google-generativeai
      系統需安裝 poppler（Ubuntu: sudo apt install poppler-utils）
"""

from __future__ import annotations

import base64
import io
import json
import os
import re
import sys


def _get_api_key() -> str | None:
    """從環境變數或 .streamlit/secrets.toml 讀取 GEMINI_API_KEY。"""
    key = os.environ.get("GEMINI_API_KEY", "").strip()
    if key:
        return key
    secrets_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".streamlit", "secrets.toml")
    if not os.path.isfile(secrets_path):
        return None
    try:
        with open(secrets_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                if k.strip() == "GEMINI_API_KEY":
                    v = v.strip()
                    if (v.startswith('"') and v.endswith('"')) or (v.startswith("'") and v.endswith("'")):
                        v = v[1:-1].replace("\\n", "\n").replace("\\t", "\t")
                    return (v or "").strip() or None
    except Exception:
        pass
    return None


def _pdf_to_images(pdf_path: str, dpi: int = 200):
    """PDF 轉成 PIL 圖片列表。"""
    from pdf2image import convert_from_path
    return convert_from_path(pdf_path, dpi=dpi)


def _extract_images_from_pdf(pdf_path: str) -> list[list[bytes]]:
    """每頁抽取的圖片 bytes 列表。images_per_page[i] = 第 i 頁的 [img1_bytes, img2_bytes, ...]。"""
    try:
        import fitz
    except ImportError:
        return []
    out = []
    doc = fitz.open(pdf_path)
    for page in doc:
        img_list = []
        for img in page.get_images():
            xref = img[0]
            base = doc.extract_image(xref)
            img_list.append(base["image"])
        out.append(img_list)
    doc.close()
    return out


def _call_gemini_vision(api_key: str, image_bytes: bytes, prompt: str, model_name: str = "gemini-2.0-flash") -> str:
    """呼叫 Gemini Vision API，回傳文字。"""
    import requests
    img_b64 = base64.b64encode(image_bytes).decode()
    m = model_name if model_name.startswith("models/") else f"models/{model_name}"
    url = f"https://generativelanguage.googleapis.com/v1beta/{m}:generateContent?key={api_key}"
    payload = {
        "contents": [{
            "parts": [
                {"text": prompt},
                {"inlineData": {"mimeType": "image/jpeg", "data": img_b64}},
            ]
        }],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 8192, "responseMimeType": "application/json"},
    }
    r = requests.post(url, json=payload, timeout=120)
    r.raise_for_status()
    data = r.json()
    if not data.get("candidates") or not data["candidates"][0].get("content", {}).get("parts"):
        raise RuntimeError("Gemini 未回傳內容")
    return data["candidates"][0]["content"]["parts"][0].get("text", "").strip()


def _parse_page_json(raw: str) -> dict:
    """從模型輸出解析 JSON（允許 ```json ... ``` 包裝）。"""
    raw = raw.strip()
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
    if m:
        raw = m.group(1).strip()
    return json.loads(raw)


def _build_docx_from_pages(
    pages_data: list[dict],
    images_per_page: list[list[bytes]],
    page_pil_images: list,
) -> bytes:
    """依 pages_data（每頁 JSON）與圖片建出 .docx bytes。"""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    for page_idx, page in enumerate(pages_data):
        blocks = page.get("blocks") or []
        page_images = images_per_page[page_idx] if page_idx < len(images_per_page) else []
        pil_page = page_pil_images[page_idx] if page_idx < len(page_pil_images) else None
        img_index = 0

        for block in blocks:
            t = (block.get("type") or "").strip().lower()
            text = (block.get("text") or "").strip()

            if t == "title":
                p = doc.add_paragraph(text)
                p.style = "Title"
            elif t == "subtitle":
                p = doc.add_paragraph(text)
                p.style = "Subtitle"
            elif t == "heading":
                level = block.get("level", 1)
                p = doc.add_paragraph(text)
                p.style = f"Heading {min(level, 3)}"
            elif t == "paragraph":
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)
            elif t == "bullet_list":
                for item in block.get("items") or []:
                    p = doc.add_paragraph(str(item).strip(), style="List Bullet")
                    p.paragraph_format.space_after = Pt(3)
            elif t == "table":
                header = block.get("header") or []
                rows = block.get("rows") or []
                if header or rows:
                    col_count = max(len(header), max((len(r) for r in rows), default=0)) or 1
                    row_count = (1 if header else 0) + len(rows)
                    table = doc.add_table(rows=row_count, cols=col_count)
                    table.style = "Table Grid"
                    r = 0
                    if header:
                        for c, cell_text in enumerate(header):
                            if c < col_count:
                                table.rows[r].cells[c].text = str(cell_text)
                        r += 1
                    for row in rows:
                        for c, cell_text in enumerate(row):
                            if c < col_count:
                                table.rows[r].cells[c].text = str(cell_text)
                        r += 1
            elif t == "image":
                if img_index < len(page_images):
                    try:
                        doc.add_picture(io.BytesIO(page_images[img_index]), width=Inches(3.0))
                    except Exception:
                        if pil_page:
                            buf = io.BytesIO()
                            pil_page.save(buf, format="PNG")
                            buf.seek(0)
                            doc.add_picture(buf, width=Inches(4.0))
                    img_index += 1
                elif pil_page:
                    buf = io.BytesIO()
                    pil_page.save(buf, format="PNG")
                    buf.seek(0)
                    doc.add_picture(buf, width=Inches(4.0))
                    img_index += 1

        if page_idx < len(pages_data) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


PROMPT_ZH = """请将这张 PDF 页面的内容解析成结构化 JSON。只输出一个 JSON 对象，不要其他说明或 markdown 标记。

格式要求：
{
  "page": 1,
  "blocks": [
    {"type": "title", "text": "主标题文字"},
    {"type": "subtitle", "text": "副标题"},
    {"type": "heading", "level": 1, "text": "标题文字"},
    {"type": "paragraph", "text": "段落内容"},
    {"type": "bullet_list", "items": ["项目1", "项目2"]},
    {"type": "table", "header": ["列1", "列2"], "rows": [["a","b"], ["c","d"]]},
    {"type": "image", "id": "img1"}
  ]
}

规则：
1. 按阅读顺序列出所有区块。
2. type 只能是：title, subtitle, heading, paragraph, bullet_list, table, image 之一。
3. 表格用 header（可选）和 rows 二维数组。
4. 图片区块用 type:"image" 和 id:"img1", "img2"... 按出现顺序编号。
5. 只输出上述 JSON，不要 ``` 包裹。"""


def run(pdf_path: str, out_path: str | None = None, api_key: str | None = None, model: str = "gemini-2.0-flash") -> str:
    """執行 PDF → Word（AI 版面還原）。回傳輸出的 .docx 路徑。"""
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"找不到 PDF：{pdf_path}")

    key = api_key or _get_api_key()
    if not key:
        raise RuntimeError("請設定 GEMINI_API_KEY（環境變數或 .streamlit/secrets.toml）")

    if out_path is None:
        base = os.path.splitext(os.path.basename(pdf_path))[0]
        out_path = os.path.join(os.path.dirname(pdf_path), f"{base}_ai_layout.docx")

    print("正在將 PDF 轉成圖片…")
    page_pil_images = _pdf_to_images(pdf_path, dpi=200)
    if not page_pil_images:
        raise RuntimeError("PDF 無有效頁面或需安裝 poppler")

    print("正在從 PDF 抽取內嵌圖片…")
    images_per_page = _extract_images_from_pdf(pdf_path)

    pages_data = []
    for i, pil_img in enumerate(page_pil_images):
        print(f"正在解析第 {i + 1}/{len(page_pil_images)} 頁…")
        if pil_img.mode != "RGB":
            pil_img = pil_img.convert("RGB")
        pil_img.thumbnail((1600, 1600))
        buf = io.BytesIO()
        pil_img.save(buf, format="JPEG", quality=88)
        img_bytes = buf.getvalue()

        raw = _call_gemini_vision(key, img_bytes, PROMPT_ZH, model)
        try:
            data = _parse_page_json(raw)
        except json.JSONDecodeError as e:
            print(f"  第 {i+1} 頁 JSON 解析失敗，改用純文字回退：{e}")
            data = {
                "page": i + 1,
                "blocks": [{"type": "paragraph", "text": raw[:5000]}]
            }
        pages_data.append(data)

    print("正在組裝 Word…")
    docx_bytes = _build_docx_from_pages(pages_data, images_per_page, page_pil_images)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with open(out_path, "wb") as f:
        f.write(docx_bytes)
    print(f"已儲存：{out_path}")
    return out_path


def main():
    if len(sys.argv) < 2:
        print("用法：python ai_pdf_to_word_prototype.py <PDF路徑> [輸出.docx路徑]")
        print("請設定 GEMINI_API_KEY（環境變數或 .streamlit/secrets.toml）")
        sys.exit(1)
    pdf_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) > 2 else None
    try:
        run(pdf_path, out_path=out_path)
    except Exception as e:
        print(f"錯誤：{e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
